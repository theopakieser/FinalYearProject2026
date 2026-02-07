# File: snapshot.py
# Description: 
# Author: Theo Pakieser
# Date: 03/02/2026

#imports
from __future__ import annotations #stores types as strings, resolves them later
from dataclasses import dataclass #imports @dataclass to autogenerate _init_ and _repr, etc
from pathlib import Path #path fields
from typing import Optional #not every file will produce a text snapshot
import re #regex module, strips space/tabs

TEXT_EXTS = { #list of file extensions treated as plain text
    ".txt", ".log", ".csv", ".json", ".xml", ".ini", ".cfg",
    ".md", ".py", ".cs", ".java", ".js", ".ts", ".html", ".css"
    }

@dataclass (frozen = True) #makes class a dataclass, immutable after creation
                            #to prevent accidental mutation of snapshots during scanning/verification
class TextSnapshot: #class definition
    text : str #extracted and normalised text
    kind : str #where text comes from

_ws_re = re.compile(r"[ \t]+$") #compiles regex pattern once
#one or more spaces or tab characters at the end of the line

def normalise_text(s : str) -> str: #make extracted text stable
    #does 3 noramlisations: windows newlines, old mac newlines and removes null bytes
    s = s.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    lines = s.plit("\n")
    #for each line run regex substitution, removes traling spaves/tabs
    #prevents hash changes caused by accidental whitespace
    lines = [_ws_re.sub("", line) for line in lines]
    #joins lines back togther and removes whitespace at the end and start of the whole text
    out = "\n".join(lines).strip()
    #enforce exactly one trailing newline, if empty, return empty string
    return out + "\n" if out else ""

def _read_text(path : Path) -> str: #reads plaintext and returns decoded string
    raw = path.read_bytes() #reads files as bytes
    try:
        return raw.decode("utf-8") #attempts decode
    except UnicodeDecodeError:#if it fails, tries other decodes
        return raw.decode("latin-1", errors="replace" )
    
def _pdf_text(path : Path) -> str: #extracts text from pdf
    import fitz #pymupdf
    parts = []
    with fitz.open(path) as doc: #opens pdfs safely and ensures closed propoerly
        for page in doc:
            parts.append(page.get_text("text")) #extracts plaintext from each page
    return "".join(parts) #combines pages into one big string

def _docx_text(path : Path) -> str: #extracts from word doc
    from docx import Document
    doc = Document(path)

    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text) #appends non-empty paragraphs

    for table in doc.tables: #loops over tables too
        for row in table.rows:
            #takes each cell's text, strips it and joins cells with tabs
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(): #only store if not empty
                parts.append(row_text)

    return "\n".join(parts) 

def extract_text_snapshot(path: Path) -> Optional[TextSnapshot]: #choose extract based on extension
    ext = path.suffix.lower() #takes files ext and makes them all lowercase

    if ext in TEXT_EXTS:
        return TextSnapshot(normalise_text(_read_text(path)), "text")
    
    if ext == ".pdf":
        return TextSnapshot(normalise_text(_pdf_text(path)), "pdf_text")
    
    if ext == ".docx":
        return TextSnapshot(normalise_text(_docx_text(path)), "docx_text")
    
    return None # unsupported file type will retun none in baseline

